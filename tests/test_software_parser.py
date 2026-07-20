import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK

from parsers import (
    certificate_key,
    normalize_certificate_number,
    parse_software_certificate_text,
    parse_software_docx,
)
from software_doc import generate_software_batch_docx, generate_software_docx


def _set_lines(cell, lines, *, soft_breaks=False):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    for index, line in enumerate(lines):
        if index and soft_breaks:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif index:
            paragraph = cell.add_paragraph()
        paragraph.add_run(line)


class SoftwareParserTest(unittest.TestCase):
    def test_finds_reordered_table_and_multiline_values(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / 'software.docx'
            doc = Document()
            doc.add_table(rows=1, cols=2).cell(0, 0).text = 'Служебная часть'

            table = doc.add_table(rows=3, cols=6)
            headers = [
                'Фамилия имя отчество',
                'Вклад авторов, %',
                'Наименование результата',
                'Должность авторов',
                'Сведения о регистрации',
                '№ п/п',
            ]
            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = header

            row = table.rows[1].cells
            _set_lines(row[0], ['Иванов И.И.', 'Петров П.П.'], soft_breaks=True)
            _set_lines(row[1], ['40%', '60%'], soft_breaks=True)
            row[2].text = 'Тестовая программа'
            _set_lines(row[3], ['Научный сотрудник', 'Инженер'], soft_breaks=True)
            _set_lines(row[4], [
                'Свидетельство о государственной регистрации программы для ЭВМ',
                'Номер свидетельства: RU-2026123456',
                'Дата государственной регистрации в реестре: 2026-07-17',
            ])
            row[5].text = '1'

            article_row = table.rows[2].cells
            article_row[0].text = 'Иванов И.И.'
            article_row[1].text = '100%'
            article_row[2].text = 'Обычная статья'
            article_row[3].text = 'Научный сотрудник'
            article_row[4].text = 'Журнал, 2026, № 1'
            article_row[5].text = '2'
            doc.save(path)

            parsed = parse_software_docx(str(path))

        self.assertEqual(len(parsed), 1)
        entry = parsed[0]
        self.assertEqual(entry['title'], 'Тестовая программа')
        self.assertEqual(entry['certificate_number'], 'RU 2026123456')
        self.assertEqual(entry['registration_date'], '17.07.2026')
        self.assertEqual(
            [author['contribution_percent'] for author in entry['authors']],
            [40, 60],
        )
        self.assertEqual(
            [author['position'] for author in entry['authors']],
            ['Научный сотрудник', 'Инженер'],
        )

    def test_generated_document_round_trip(self):
        data = {
            'title': 'Система контроля',
            'certificate_number': '2026987654',
            'registration_date': '2026-07-17',
            'output_data': '',
            'authors': [
                {
                    'full_name': 'Иванов И.И.',
                    'position': 'Научный сотрудник',
                    'contribution_percent': 100,
                }
            ],
        }
        content = generate_software_docx(data, {})

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / 'generated.docx'
            path.write_bytes(content)
            parsed = parse_software_docx(str(path))

        self.assertEqual(len(parsed), 1)
        self.assertEqual(parsed[0]['certificate_number'], 'RU 2026987654')
        self.assertEqual(parsed[0]['registration_date'], '17.07.2026')
        self.assertEqual(parsed[0]['authors'][0]['contribution_percent'], 100)

    def test_certificate_normalization_and_comparison(self):
        variants = ('RU 2026123456', 'RU2026123456', 'RU-2026123456', '2026123456')
        self.assertEqual(
            {normalize_certificate_number(value) for value in variants},
            {'RU 2026123456'},
        )
        self.assertEqual({certificate_key(value) for value in variants}, {'2026123456'})

    def test_parses_pdf_text_and_ignores_application_number(self):
        text = """№ 2026680579
Модуль потокового вывода ответов
Правообладатель: Тихонов Сергей Сергеевич (RU)
Авторы: Тихонов Сергей Сергеевич (RU), Григоренко
Александр Георгиевич (RU), Васильев Никита Алексеевич (RU)
Заявка № 2026667857
Дата поступления 26 мая 2026 г.
Дата государственной регистрации
в Реестре программ для ЭВМ 10 июля 2026 г.
"""

        parsed = parse_software_certificate_text(text)

        self.assertEqual(parsed['certificate_number'], 'RU 2026680579')
        self.assertNotIn('2026667857', parsed['output_data'])
        self.assertEqual(parsed['registration_date'], '10.07.2026')
        self.assertEqual(parsed['title'], 'Модуль потокового вывода ответов')
        self.assertEqual(len(parsed['authors']), 3)
        self.assertEqual(
            [author['contribution_percent'] for author in parsed['authors']],
            [0, 0, 0],
        )

    def test_generated_batch_document_round_trip(self):
        entries = []
        for index in range(2):
            entries.append({
                'title': f'Программа {index + 1}',
                'certificate_number': f'RU 202688800{index + 1}',
                'registration_date': '10.07.2026',
                'output_data': '',
                'authors': [{
                    'full_name': 'Иванов И.И.',
                    'position': 'Научный сотрудник',
                    'contribution_percent': 100,
                }],
            })
        content = generate_software_batch_docx(entries, {})

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / 'batch.docx'
            path.write_bytes(content)
            parsed = parse_software_docx(str(path))

        self.assertEqual(len(parsed), 2)
        self.assertEqual(
            [entry['certificate_number'] for entry in parsed],
            ['RU 2026888001', 'RU 2026888002'],
        )


if __name__ == '__main__':
    unittest.main()
