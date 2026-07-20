"""Generate 'Докладная записка ПО' docx matching the official template format."""
import io
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_paragraph(cell, text, bold=False, center=False, size=10):
    p = cell.paragraphs[0]
    p.clear()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    return p


def _add_cell_lines(cell, lines, bold=False, size=10, center=False):
    """Add multiple lines as separate paragraphs inside a cell."""
    cell.paragraphs[0].clear()
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = align
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1


def _set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement('w:tblHeader')
    repeat.set(qn('w:val'), 'true')
    tr_pr.append(repeat)


def _prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement('w:cantSplit'))


def generate_software_batch_docx(sw_entries: list[dict], profile: dict) -> bytes:
    """
    sw_entries: [{ title, certificate_number, registration_date, output_data,
                   authors: [{full_name, position, contribution_percent}] }]
    profile: { last_name, first_patronymic, position, unit }
    """
    if not sw_entries:
        raise ValueError('Для докладной записки не передано ни одной программы')
    doc = Document()

    # ── Page: A4 landscape ───────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(2.25)
    from docx.enum.section import WD_ORIENT
    section.orientation = WD_ORIENT.LANDSCAPE

    # ── Header paragraphs (right-aligned, 12pt) ───────────────────────────────
    def add_right(text, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = bold
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(0)
        return p

    add_right('Приложение к докладной записке')
    add_right('результаты работы в области научных исследований')

    doc.add_paragraph()  # spacer

    # ── Table ────────────────────────────────────────────────────────────────
    # Col widths (EMU): №п/п | Название | Выходные данные | ФИО | Должность | Вклад%
    col_widths_cm = [1.1, 5.4, 5.4, 3.4, 5.8, 2.7]

    table = doc.add_table(rows=len(sw_entries) + 1, cols=6)
    table.style = 'Table Grid'
    table.autofit = False
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    table._tbl.tblPr.append(layout)

    # Fix column widths
    for ci, width in enumerate(col_widths_cm):
        table.columns[ci].width = Cm(width)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = Cm(col_widths_cm[ci])
            _set_cell_borders(cell)

    body_size = 8 if len(sw_entries) >= 4 else (9 if len(sw_entries) >= 3 else 10)
    header_size = 8.5 if len(sw_entries) >= 4 else 10

    # Header row
    hdrs = [
        '№\nп/п',
        'Название статьи',
        'Выходные данные статьи',
        'Фамилия имя отчество',
        'Должность авторов с указанием кафедры/\nуправления/отдела',
        'Вклад авторов в статью/\nмонографию/\nпатент\n(в сумме 100%)',
    ]
    for ci, hdr in enumerate(hdrs):
        cell = table.rows[0].cells[ci]
        _cell_paragraph(cell, hdr, bold=True, center=True, size=header_size)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_repeat_header(table.rows[0])

    for row_index, sw_data in enumerate(sw_entries, start=1):
        authors = sw_data.get('authors', [])
        fio_lines = [a.get('full_name', '') for a in authors]
        pos_lines = [a.get('position', '') for a in authors]
        pct_lines = [f"{a.get('contribution_percent', 0)}%" for a in authors]

        cert = sw_data.get('certificate_number', '')
        rdate = sw_data.get('registration_date', '')
        out = sw_data.get('output_data', '')
        if cert and rdate:
            output_text = (
                "Тип: свидетельство о государственной регистрации программы для ЭВМ\n"
                f"Номер свидетельства: {cert}\n"
                f"Дата государственной регистрации в реестре программ для ЭВМ: {rdate}"
            )
        else:
            output_text = out

        data_row = table.rows[row_index]
        _prevent_row_split(data_row)
        _cell_paragraph(data_row.cells[0], str(row_index), center=True, size=body_size)
        _cell_paragraph(data_row.cells[1], sw_data.get('title', ''), size=body_size)
        _add_cell_lines(data_row.cells[2], output_text.split('\n'), size=body_size)
        _add_cell_lines(data_row.cells[3], fio_lines, size=body_size)
        _add_cell_lines(data_row.cells[4], pos_lines, size=body_size)
        _add_cell_lines(data_row.cells[5], pct_lines, center=False, size=body_size)

        for cell in data_row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_software_docx(sw_data: dict, profile: dict) -> bytes:
    return generate_software_batch_docx([sw_data], profile)
