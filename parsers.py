import re
from io import BytesIO
from docx import Document


def _normalize(text: str) -> str:
    text = (text or '').replace('\xa0', ' ')
    text = re.sub(r'(\d)\s*\.\s*(\d)', r'\1.\2', text)
    text = re.sub(r'№\s+', '№', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def normalize_certificate_number(value: str) -> str:
    """Return one stable representation for standard Russian certificate numbers."""
    text = _normalize(value)
    match = re.fullmatch(r'(?:RU[\s\-]*)?(\d{6,})', text, re.IGNORECASE)
    if match:
        return f"RU {match.group(1)}"
    return text


def certificate_key(value: str) -> str:
    """Comparison key that treats RU 123, RU123 and 123 as the same number."""
    normalized = normalize_certificate_number(value)
    match = re.fullmatch(r'RU\s+(\d+)', normalized, re.IGNORECASE)
    if match:
        return match.group(1)
    return re.sub(r'[^A-ZА-ЯЁ0-9]', '', normalized.upper())


def _normalize_registration_date(value: str) -> str:
    value = (value or '').strip()
    match = re.fullmatch(r'(\d{4})[-/.](\d{2})[-/.](\d{2})', value)
    if match:
        return f"{match.group(3)}.{match.group(2)}.{match.group(1)}"
    match = re.fullmatch(r'(\d{2})[-/.](\d{2})[-/.](\d{4})', value)
    if match:
        return f"{match.group(1)}.{match.group(2)}.{match.group(3)}"
    return value


def parse_order_docx(file_path: str) -> dict:
    doc = Document(file_path)
    full_text = ' '.join(p.text for p in doc.paragraphs)
    full_text = _normalize(full_text)

    result = {
        'level': 'academy',
        'number': '',
        'title': '',
        'order_date': '',
        'deadline_type': 'monthly',
        'deadline_date': None,
        'workload_hours': None,
    }

    text_lower = full_text.lower()

    # Level detection
    if any(kw in text_lower for kw in ['вышестоящего органа', 'главного управления', 'министерства обороны']):
        result['level'] = 'higher'
    if any(kw in text_lower for kw in ['руководства академии', 'начальника вас', 'руководящего состава военной академии']):
        result['level'] = 'academy'

    # Order number + date
    m = re.search(r'(?:Приказани[ея]|Приказ)[^№]*№(\d+)[^\d]*от\s+([\d.]+)', full_text)
    if m:
        result['number'] = m.group(1)
        result['order_date'] = m.group(2)

    # Theme
    m = re.search(r'Тема\s*:\s*[«\"](.+?)[»\"]', full_text)
    if m:
        result['title'] = m.group(1).strip()

    # Deadline
    if 'ежемесячно' in text_lower:
        result['deadline_type'] = 'monthly'
    else:
        m = re.search(r'Дата выполнения\s*:\s*([\d.]+)', full_text)
        if m:
            result['deadline_type'] = 'date'
            result['deadline_date'] = m.group(1)

    # Workload
    m = re.search(r'Трудозатраты\s*:\s*(\d+)\s*час', full_text)
    if m:
        result['workload_hours'] = int(m.group(1))

    return result


def _cell_paragraphs(cell) -> list[str]:
    """Return non-empty lines from a cell.
    Handles both real paragraph breaks and soft line breaks (<w:br/>),
    since cell.text uses '\\n' for both.
    """
    lines = []
    for p in cell.paragraphs:
        for line in p.text.split('\n'):
            line = line.replace('\xa0', ' ').strip().rstrip(',').strip()
            if line:
                lines.append(line)
    return lines


_DEFAULT_SOFTWARE_COLUMNS = {
    'title': 1,
    'output': 2,
    'authors': 3,
    'positions': 4,
    'contribution': 5,
}


def _find_header_column(headers: list[str], *needles: str) -> int | None:
    for index, header in enumerate(headers):
        if any(needle in header for needle in needles):
            return index
    return None


def _software_table_layout(table) -> tuple[int, dict[str, int]] | None:
    """Find the header and columns, falling back to the legacy six-column layout."""
    for row_index, row in enumerate(table.rows[:3]):
        headers = [_normalize(cell.text).lower() for cell in row.cells]
        columns = {
            'title': _find_header_column(headers, 'название', 'наименование'),
            'output': _find_header_column(headers, 'выходн', 'сведения о регистрац'),
            'authors': _find_header_column(headers, 'фамили', 'фио'),
            'positions': _find_header_column(headers, 'должност'),
            'contribution': _find_header_column(headers, 'вклад', 'доля', '%'),
        }
        if all(index is not None for index in columns.values()):
            return row_index, columns

    if table.rows and len(table.rows[0].cells) >= 6:
        return 0, dict(_DEFAULT_SOFTWARE_COLUMNS)
    return None


def _looks_like_software_output(value: str) -> bool:
    text = _normalize(value).lower()
    if not re.search(r'свидетельств\w*', text):
        return False
    return bool(
        re.search(r'регистрац\w*', text)
        or 'номер свидетельства' in text
        or re.search(r'№\s*(?:ru[\s\-]*)?\d{6,}', text, re.IGNORECASE)
        or re.search(r'\bru[\s\-]*\d{6,}\b', text, re.IGNORECASE)
    )


def _parse_software_row(cells, columns: dict[str, int] | None = None) -> dict:
    """Parse one data row from the software table."""
    columns = columns or _DEFAULT_SOFTWARE_COLUMNS
    entry = {
        'title': '',
        'certificate_number': '',
        'registration_date': '',
        'output_data': '',
        'authors': [],
    }

    entry['title'] = _normalize(cells[columns['title']].text)

    output_text = _normalize(cells[columns['output']].text)
    entry['output_data'] = output_text

    certificate_patterns = (
        r'Номер\s+свидетельства\s*[:№]?\s*((?:RU[\s\-]*)?\d{6,})',
        r'свидетельств\w*[^№]{0,160}№\s*((?:RU[\s\-]*)?\d{6,})',
        r'регистрац\w*[^№]{0,160}№\s*((?:RU[\s\-]*)?\d{6,})',
        r'\b(RU[\s\-]*\d{6,})\b',
    )
    for pattern in certificate_patterns:
        match = re.search(pattern, output_text, re.IGNORECASE)
        if match:
            entry['certificate_number'] = normalize_certificate_number(match.group(1))
            break

    date_patterns = (
        r'Дата\s+государственной\s+регистрации[^:]*:\s*(\d{2}[./-]\d{2}[./-]\d{4}|\d{4}[./-]\d{2}[./-]\d{2})',
        r'Дата\s+регистрации[^:]*:\s*(\d{2}[./-]\d{2}[./-]\d{4}|\d{4}[./-]\d{2}[./-]\d{2})',
        r'опубл\.?\s*(\d{2}[./-]\d{2}[./-]\d{4}|\d{4}[./-]\d{2}[./-]\d{2})',
    )
    for pattern in date_patterns:
        match = re.search(pattern, output_text, re.IGNORECASE)
        if match:
            entry['registration_date'] = _normalize_registration_date(match.group(1))
            break

    fio_list = _cell_paragraphs(cells[columns['authors']])
    pos_list = _cell_paragraphs(cells[columns['positions']])

    pct_list = []
    for line in _cell_paragraphs(cells[columns['contribution']]):
        pct_list.extend(int(value) for value in re.findall(r'(\d+)\s*%', line))

    for i, fio in enumerate(fio_list):
        entry['authors'].append({
            'full_name': fio,
            'position': pos_list[i] if i < len(pos_list) else '',
            'contribution_percent': pct_list[i] if i < len(pct_list) else 0,
            'points_claimed': 0,
        })

    return entry


def _parse_article_row(cells) -> dict:
    """Parse one data row from the research-output table as an article entry."""
    entry = {'title': '', 'publication': '', 'authors': []}
    entry['title'] = _normalize(cells[1].text.strip())
    entry['publication'] = _normalize(cells[2].text.strip())

    fio_list = _cell_paragraphs(cells[3])
    pos_list = _cell_paragraphs(cells[4])

    pct_list = []
    for p in cells[5].paragraphs:
        full = ''.join(r.text for r in p.runs).strip() or p.text.strip()
        full = _normalize(full)
        m = re.search(r'(\d+)\s*%', full)
        if m:
            pct_list.append(int(m.group(1)))

    for i, fio in enumerate(fio_list):
        entry['authors'].append({
            'full_name': fio,
            'position': pos_list[i] if i < len(pos_list) else '',
            'contribution_percent': pct_list[i] if i < len(pct_list) else 0,
        })
    return entry


def parse_article_docx(file_path: str) -> list:
    """Returns article entries from a research-output dokkladnaya.
    Rows where the output column contains a software certificate are skipped.
    """
    doc = Document(file_path)
    if not doc.tables:
        return []

    results = []
    for row in doc.tables[0].rows[1:]:
        cells = row.cells
        if len(cells) < 6:
            continue
        title = cells[1].text.strip()
        if not title:
            continue
        output = cells[2].text.strip()
        if _looks_like_software_output(output):
            continue  # ПО entry — belongs to software parser
        results.append(_parse_article_row(cells))
    return results


def parse_software_docx(file_path: str) -> list:
    """Returns a list of software entries (one per table data row).
    Rows without a software certificate in the output column are skipped.
    """
    doc = Document(file_path)
    if not doc.tables:
        return []

    results = []
    for table in doc.tables:
        layout = _software_table_layout(table)
        if not layout:
            continue
        header_index, columns = layout
        required_index = max(columns.values())
        for row in table.rows[header_index + 1:]:
            cells = row.cells
            if len(cells) <= required_index:
                continue
            title = cells[columns['title']].text.strip()
            if not title:
                continue
            output = cells[columns['output']].text
            if not _looks_like_software_output(output):
                continue  # article entry — skip
            results.append(_parse_software_row(cells, columns))

    return results


_RU_MONTHS = {
    'января': 1,
    'февраля': 2,
    'марта': 3,
    'апреля': 4,
    'мая': 5,
    'июня': 6,
    'июля': 7,
    'августа': 8,
    'сентября': 9,
    'октября': 10,
    'ноября': 11,
    'декабря': 12,
}


def parse_software_certificate_text(text: str) -> dict:
    """Parse a Rospatent software certificate without using its application number."""
    text = (text or '').replace('\xa0', ' ')
    if not text.strip():
        raise ValueError('В PDF отсутствует текстовый слой')

    certificate_match = re.search(r'(?m)^\s*№\s*(\d{7,})\s*$', text)
    if not certificate_match:
        raise ValueError('Не найден номер свидетельства')
    certificate = normalize_certificate_number(certificate_match.group(1))

    holder_match = re.search(r'Правообладатель\s*:', text, re.IGNORECASE)
    if not holder_match or holder_match.start() <= certificate_match.end():
        raise ValueError('Не найдено название программы')
    title = _normalize(text[certificate_match.end():holder_match.start()])
    if not title:
        raise ValueError('Не найдено название программы')

    authors_match = re.search(
        r'Авторы\s*:\s*(.*?)\s*Заявка\s*№', text, re.IGNORECASE | re.DOTALL
    )
    if not authors_match:
        raise ValueError('Не найден список авторов')
    authors_text = re.sub(r'\(\s*RU\s*\)', '', authors_match.group(1), flags=re.IGNORECASE)
    author_names = [_normalize(name) for name in authors_text.split(',')]
    author_names = [name for name in author_names if name]
    if not author_names:
        raise ValueError('Не найден список авторов')

    registration_match = re.search(
        r'Дата\s+государственной\s+регистрации.{0,180}?'
        r'(\d{1,2})\s+'
        r'(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+'
        r'(\d{4})\s*г?\.?',
        text,
        re.IGNORECASE | re.DOTALL,
    )
    if not registration_match:
        raise ValueError('Не найдена дата государственной регистрации')
    day = int(registration_match.group(1))
    month = _RU_MONTHS[registration_match.group(2).lower()]
    year = int(registration_match.group(3))
    registration_date = f'{day:02d}.{month:02d}.{year:04d}'

    authors = []
    for full_name in author_names:
        authors.append({
            'full_name': full_name,
            'position': '',
            # В свидетельстве вклад авторов не указан. Его обязательно задаёт
            # пользователь после распознавания, поэтому ничего не вычисляем.
            'contribution_percent': 0,
            'points_claimed': 0,
        })

    output_data = (
        'Тип: свидетельство о государственной регистрации программы для ЭВМ\n'
        f'Номер свидетельства: {certificate}\n'
        'Дата государственной регистрации в реестре программ для ЭВМ: '
        f'{registration_date}'
    )
    return {
        'title': title,
        'certificate_number': certificate,
        'registration_date': registration_date,
        'output_data': output_data,
        'authors': authors,
    }


def parse_software_certificate_pdf(source: bytes) -> dict:
    """Extract text from a PDF certificate and parse the fields needed by the memo."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError('Для чтения PDF не установлен пакет pypdf') from exc

    reader = PdfReader(BytesIO(source))
    text = '\n'.join((page.extract_text() or '') for page in reader.pages)
    return parse_software_certificate_text(text)
